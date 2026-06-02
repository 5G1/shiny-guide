from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
from io import BytesIO
from docx import Document
from PIL import Image
import pandas as pd
import pytesseract
import fitz
import tempfile
import re
import os


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


WINDOWS_TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

if os.path.exists(WINDOWS_TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = WINDOWS_TESSERACT_PATH


DEVELOPER_NAME = "스프링툴바"
DEVELOPER_CONTACT = "springtoolbar@gmail.com"
MAX_FILES = 10


@app.get("/")
def home():
    return {
        "message": "Springtool Backend Running",
        "max_files": MAX_FILES,
        "developer": DEVELOPER_NAME,
        "contact": DEVELOPER_CONTACT,
    }


def split_sentences(text: str):
    sentences = re.split(r"(?<=[。！？.!?])", text)
    return [s.strip() for s in sentences if s.strip()]


def clean_keywords(keyword_input: str):
    keywords = re.split(r"[,，\n]", keyword_input)
    return [k.strip() for k in keywords if k.strip()]


def clean_text_for_count(text: str):
    return text.replace(" ", "").replace("\n", "").replace("\t", "")


def ocr_page_with_pymupdf(page):
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        return text
    except Exception:
        return ""


def analyze_pdf_file(
    file_name: str,
    pdf_path: str,
    keywords: List[str],
    use_ocr: bool,
    show_context: bool,
):
    doc = fitz.open(pdf_path)

    file_results = []
    file_logs = []
    file_total_text = ""

    keyword_stats = {
        keyword: {
            "count": 0,
            "keyword_length": len(keyword),
        }
        for keyword in keywords
    }

    for page_index in range(len(doc)):
        page = doc[page_index]
        page_number = page_index + 1

        text = page.get_text()
        extraction_method = "Text extraction"

        if not text.strip() and use_ocr:
            text = ocr_page_with_pymupdf(page)
            extraction_method = "OCR" if text.strip() else "OCR unavailable or no text found"
        elif not text.strip():
            extraction_method = "No text"

        file_total_text += text

        file_logs.append(
            {
                "pdf_name": file_name,
                "page_number": page_number,
                "method": extraction_method,
                "text_length": len(clean_text_for_count(text)),
            }
        )

        sentences = split_sentences(text)

        for keyword in keywords:
            keyword_stats[keyword]["count"] += text.count(keyword)

        for sentence_index, sentence in enumerate(sentences):
            for keyword in keywords:
                if keyword in sentence:
                    if show_context:
                        previous_sentence = sentences[sentence_index - 1] if sentence_index > 0 else ""
                        next_sentence = (
                            sentences[sentence_index + 1]
                            if sentence_index < len(sentences) - 1
                            else ""
                        )
                        context = f"{previous_sentence} {sentence} {next_sentence}".strip()
                    else:
                        context = sentence

                    file_results.append(
                        {
                            "pdf_name": file_name,
                            "keyword": keyword,
                            "page_number": page_number,
                            "matched_sentence": sentence,
                            "context": context,
                            "method": extraction_method,
                        }
                    )

    file_total_chars = len(clean_text_for_count(file_total_text))

    file_stats = []
    for keyword, stat in keyword_stats.items():
        total_keyword_chars = stat["count"] * stat["keyword_length"]
        ratio = (
            total_keyword_chars / file_total_chars * 100
            if file_total_chars > 0
            else 0
        )

        file_stats.append(
            {
                "pdf_name": file_name,
                "keyword": keyword,
                "count": stat["count"],
                "keyword_length": stat["keyword_length"],
                "total_keyword_chars": total_keyword_chars,
                "pdf_total_chars": file_total_chars,
                "ratio": f"{ratio:.4f}%",
            }
        )

    return {
        "pdf_name": file_name,
        "pages": len(doc),
        "total_results": len(file_results),
        "total_chars": file_total_chars,
        "stats": file_stats,
        "results": file_results,
        "logs": file_logs,
    }


@app.post("/search")
async def search_pdf(
    files: List[UploadFile] = File(...),
    keywords: str = Form(...),
    use_ocr: str = Form("false"),
    show_context: str = Form("false"),
):
    if not files:
        raise HTTPException(status_code=400, detail="At least one PDF file is required.")

    if len(files) > MAX_FILES:
        raise HTTPException(
            status_code=400,
            detail=f"Maximum {MAX_FILES} PDF files can be uploaded at once.",
        )

    for file in files:
        if not file.filename.lower().endswith(".pdf"):
            raise HTTPException(
                status_code=400,
                detail=f"Only PDF files are supported: {file.filename}",
            )

    keyword_list = clean_keywords(keywords)

    if not keyword_list:
        raise HTTPException(status_code=400, detail="At least one keyword is required.")

    use_ocr_bool = use_ocr == "true"
    show_context_bool = show_context == "true"

    all_results = []
    all_stats = []
    all_logs = []
    file_summaries = []

    global_keyword_stats = {
        keyword: {
            "count": 0,
            "keyword_length": len(keyword),
            "total_keyword_chars": 0,
        }
        for keyword in keyword_list
    }

    global_total_chars = 0
    total_pages = 0

    for file in files:
        contents = await file.read()

        if not contents:
            raise HTTPException(
                status_code=400,
                detail=f"Uploaded file is empty: {file.filename}",
            )

        temp_path = None

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp:
                temp.write(contents)
                temp_path = temp.name

            analyzed = analyze_pdf_file(
                file_name=file.filename,
                pdf_path=temp_path,
                keywords=keyword_list,
                use_ocr=use_ocr_bool,
                show_context=show_context_bool,
            )

            file_summaries.append(
                {
                    "pdf_name": analyzed["pdf_name"],
                    "pages": analyzed["pages"],
                    "total_results": analyzed["total_results"],
                    "total_chars": analyzed["total_chars"],
                }
            )

            total_pages += analyzed["pages"]
            global_total_chars += analyzed["total_chars"]
            all_results.extend(analyzed["results"])
            all_stats.extend(analyzed["stats"])
            all_logs.extend(analyzed["logs"])

            for stat in analyzed["stats"]:
                keyword = stat["keyword"]
                global_keyword_stats[keyword]["count"] += stat["count"]
                global_keyword_stats[keyword]["total_keyword_chars"] += stat[
                    "total_keyword_chars"
                ]

        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass

    global_stats = []
    for keyword, stat in global_keyword_stats.items():
        ratio = (
            stat["total_keyword_chars"] / global_total_chars * 100
            if global_total_chars > 0
            else 0
        )

        global_stats.append(
            {
                "pdf_name": "ALL PDF FILES",
                "keyword": keyword,
                "count": stat["count"],
                "keyword_length": stat["keyword_length"],
                "total_keyword_chars": stat["total_keyword_chars"],
                "pdf_total_chars": global_total_chars,
                "ratio": f"{ratio:.4f}%",
            }
        )

    return {
        "pdf_names": [file.filename for file in files],
        "total_files": len(files),
        "total_pages": total_pages,
        "keywords": keyword_list,
        "total_results": len(all_results),
        "file_summaries": file_summaries,
        "stats": global_stats,
        "file_stats": all_stats,
        "results": all_results,
        "logs": all_logs,
        "developer": DEVELOPER_NAME,
        "contact": DEVELOPER_CONTACT,
    }


class ExportStat(BaseModel):
    pdf_name: str
    keyword: str
    count: int
    keyword_length: int
    total_keyword_chars: int
    pdf_total_chars: int
    ratio: str


class ExportResult(BaseModel):
    pdf_name: str
    keyword: str
    page_number: int
    matched_sentence: str
    context: str
    method: str


class ExportFileSummary(BaseModel):
    pdf_name: str
    pages: int
    total_results: int
    total_chars: int


class ExportPayload(BaseModel):
    pdf_names: List[str]
    total_files: int
    total_pages: int
    keywords: List[str]
    total_results: int
    file_summaries: List[ExportFileSummary]
    stats: List[ExportStat]
    file_stats: List[ExportStat]
    results: List[ExportResult]
    developer: str = DEVELOPER_NAME
    contact: str = DEVELOPER_CONTACT


@app.post("/export/excel")
async def export_excel(payload: ExportPayload):
    output = BytesIO()

    summary_df = pd.DataFrame(
        [
            {"Item": "Total PDF Files", "Value": payload.total_files},
            {"Item": "Total Pages", "Value": payload.total_pages},
            {"Item": "Matched Results", "Value": payload.total_results},
            {"Item": "Keywords", "Value": ", ".join(payload.keywords)},
            {"Item": "Operator / Developer", "Value": DEVELOPER_NAME},
            {"Item": "Contact", "Value": DEVELOPER_CONTACT},
        ]
    )

    file_summary_df = pd.DataFrame([item.dict() for item in payload.file_summaries])
    global_stats_df = pd.DataFrame([item.dict() for item in payload.stats])
    file_stats_df = pd.DataFrame([item.dict() for item in payload.file_stats])

    results_df = pd.DataFrame(
        [
            {
                "PDF Name": item.pdf_name,
                "Keyword": item.keyword,
                "Page Number": item.page_number,
                "Matched Sentence": item.matched_sentence,
                "Context": item.context,
                "Extraction Method": item.method,
            }
            for item in payload.results
        ]
    )

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        summary_df.to_excel(writer, index=False, sheet_name="Summary")
        file_summary_df.to_excel(writer, index=False, sheet_name="File Summary")
        global_stats_df.to_excel(writer, index=False, sheet_name="Global Statistics")
        file_stats_df.to_excel(writer, index=False, sheet_name="File Statistics")
        results_df.to_excel(writer, index=False, sheet_name="Matched Sentences")

    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": 'attachment; filename="springtool_results.xlsx"'
        },
    )


@app.post("/export/word")
async def export_word(payload: ExportPayload):
    document = Document()

    document.add_heading("springtool PDF Keyword Report", level=1)

    document.add_paragraph(f"Total PDF Files: {payload.total_files}")
    document.add_paragraph(f"Total Pages: {payload.total_pages}")
    document.add_paragraph(f"Matched Results: {payload.total_results}")
    document.add_paragraph(f"Keywords: {', '.join(payload.keywords)}")
    document.add_paragraph(f"Operator / Developer: {DEVELOPER_NAME}")
    document.add_paragraph(f"Contact: {DEVELOPER_CONTACT}")

    document.add_heading("File Summary", level=2)
    for item in payload.file_summaries:
        document.add_paragraph(
            f"PDF Name: {item.pdf_name} | Pages: {item.pages} | "
            f"Results: {item.total_results} | Characters: {item.total_chars}"
        )

    document.add_heading("Global Keyword Statistics", level=2)
    for stat in payload.stats:
        document.add_paragraph(
            f"Keyword: {stat.keyword} | Count: {stat.count} | "
            f"Keyword Characters: {stat.total_keyword_chars} | "
            f"Total Characters: {stat.pdf_total_chars} | Ratio: {stat.ratio}"
        )

    document.add_heading("Matched Sentences", level=2)
    for index, item in enumerate(payload.results, start=1):
        document.add_paragraph(
            f"{index}. PDF Name: {item.pdf_name} | "
            f"Keyword: {item.keyword} | "
            f"Page Number: {item.page_number} | "
            f"Method: {item.method}"
        )
        document.add_paragraph(f"Matched Sentence: {item.matched_sentence}")
        document.add_paragraph(f"Context: {item.context}")
        document.add_paragraph("")

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="springtool_results.docx"'
        },
    )